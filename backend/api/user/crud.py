from datetime import datetime
from typing import List, Optional

import string
import secrets

from fastapi import HTTPException, status

from sqlalchemy import select, update, delete
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
from starlette.responses import FileResponse

from backend.api.auth.auth import get_password_hash
from backend.api.user.models import UserSchema, CreateUserSchema, UserParamSchema
from backend.database.tables import User, Group, Student



async def create_user(user: CreateUserSchema, session: AsyncSession) -> UserSchema:
    """
    Создает нового пользователя с хешированием пароля.

    Проверяет уникальность логина, генерирует соль и хеш пароля.

    Args:
        user: Данные для создания пользователя
        session: Асинхронная сессия БД

    Returns:
        UserSchema: Созданный пользователь

    Raises:
        HTTPException: 409 - Логин занят
    """
    # Проверка уникальности логина
    existing_user = await session.execute(select(User).where(User.login == user.login))

    if existing_user.scalars().first():
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail=f"Логин '{user.login}' уже занят"
        )

    # Генерация криптографически безопасной соли
    salt = ''.join(secrets.choice(string.ascii_letters + string.digits + string.punctuation) for _ in range(16))
    hashed_password = get_password_hash(user.password, salt)

    new_user = User(
        login=user.login,
        password=hashed_password,
        salt=salt,
        name=user.name,
        surname=user.surname,
        lastname=user.lastname,
        privilege=user.privilege
    )

    session.add(new_user)
    await session.commit()
    await session.refresh(new_user)
    return UserSchema.model_validate(new_user)

async def get_users_by_id(user_id: Optional[int] = None, session: AsyncSession = None) -> List[UserSchema]:
    """
    Получает пользователей по ID или всех пользователей.

    Args:
        user_id: ID пользователя или None для всех
        session: Асинхронная сессия

    Returns:
        Список пользователей
    """
    query = select(User)
    if user_id is not None:
        query = query.where(User.id == user_id)

    result = await session.execute(query.order_by(User.id))
    return [UserSchema.model_validate(u) for u in result.scalars().all()]

async def update_user(updates: List[UserParamSchema], session: AsyncSession, current_user: dict) -> dict:
    """
    Пакетное обновление пользователей с проверкой прав.

    Args:
        updates: Список обновлений
        session: Сессия БД
        current_user: Текущий пользователь

    Returns:
        Статус операции

    Raises:
        HTTPException: 403 - Недостаточно прав
        HTTPException: 404 - Пользователь не найден
    """
    # Предварительная проверка существования пользователей
    _ = await session.execute(select(User.id))
    existing_ids = {u for u in _.scalars().all()}
    update_ids = {u.id for u in updates}
    missing = update_ids - existing_ids
    if missing:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Пользователи {missing} не найдены"
        )

    # Пакетное обновление
    for user_update in updates:
        # Проверка прав на изменение
        if current_user['id'] != user_update.id and current_user['privilege'] < 2:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Недостаточно прав для изменения чужих данных"
            )

        # Проверка прав на изменение привилегий
        if user_update.privilege is not None and current_user['privilege'] < 2:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Только администратор может менять привилегии"
            )

        # Обновление логина
        if user_update.login:
            if current_user['id'] != user_update.id and current_user['privilege'] < 2:
                raise HTTPException(
                    status_code=status.HTTP_403_FORBIDDEN,
                    detail="Только владелец или администратор может менять логин"
                )

        # Сборка параметров обновления
        update_data = user_update.model_dump(exclude_unset=True)
        if 'password' in update_data:
            raise NotImplementedError("Смена пароля требует отдельной реализации")

        await session.execute(update(User).where(User.id == user_update.id).values(**update_data))

    await session.commit()
    return {"status": "success", "updated": len(updates)}

async def delete_user(user_ids: List[int], session: AsyncSession, current_user: dict) -> dict:
    """
    Удаляет пользователей по ID.

    Args:
        user_ids: Список ID пользователей
        session: Сессия БД
        current_user: Текущий пользователь

    Returns:
        Статус операции

    Raises:
        HTTPException: 403 - Недостаточно прав
        HTTPException: 404 - Пользователи не найдены
    """
    if current_user['privilege'] < 2:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Только администратор может удалять пользователей"
        )

    _ = await session.execute(select(User.id))
    existing = set(_.scalars().all())
    to_delete = set(user_ids)
    missing = to_delete - existing

    if missing:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Пользователи {missing} не найдены"
        )

    await session.execute(
        delete(User)
        .where(User.id.in_(user_ids))
    )
    await session.commit()
    return {"status": "success", "deleted": len(user_ids)}

async def get_group_report(group_id, session, current_user):
    group = await session.execute(select(Group).where(Group.id == group_id))
    group = group.scalar()

    if not group:
        raise HTTPException(status_code=404, detail="Group not found")

    students = await session.execute(
        select(Student)
        .where(Student.group_id == group_id)
        .options(selectinload(Student.group))
    )
    students = students.scalars().all()

    if not students:
        raise HTTPException(status_code=404, detail="No students in group")

    doc = Document()
    filename = f"Otchet_Gruppa_{datetime.today().strftime('%d.%m.%Y')}.docx"

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('ОТЧЕТ\nпо успеваемости студентов')
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title.add_run('\nРТУ МИРЭА Кафедра №250')
    title.runs[1].font.size = Pt(11.5)
    title.runs[1].font.italic = True

    doc.add_paragraph().add_run().add_break()

    info = doc.add_paragraph()
    info.add_run('Факультет: ')
    info.add_run('   Программная инженерия   \n').font.underline = True
    info.add_run('Направление подготовки: ')
    info.add_run('   Вычислительная техника   \n').font.underline = True
    info.add_run('Группа: ')
    info.add_run(f'   {group.name}   ').font.underline = True
    info.add_run('\t\t\t\t\tКурс: ___\n')

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_para.add_run(f'Дата составления: {datetime.today().strftime("%d.%m.%Y")}')
    date_para.runs[0].font.italic = True

    doc.add_paragraph().add_run().add_break()

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    columns = [
        ('№ п/п', 0.3),
        ('ФИО студента', 2.5),
        ('Тема курсовой работы', 3),
        ('Оценка', 0.7),
        ('Допуск к защите', 1)
    ]

    hdr_cells = table.rows[0].cells
    for i, (text, width) in enumerate(columns):
        hdr_cells[i].text = text
        hdr_cells[i].width = Inches(width)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True

        tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
        shade = OxmlElement('w:shd')
        shade.set(qn('w:fill'), 'E6E6E6')
        tcPr.append(shade)

    for idx, student in enumerate(students, 1):
        row_cells = table.add_row().cells

        # Номер п/п
        row_cells[0].text = str(student.educational_id)
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ФИО
        row_cells[1].text = f"{student.surname} {student.name} {student.lastname or ''}"

        # Тема курсовой
        row_cells[2].text = student.diploma.title if student.diploma else 'Не определена'

        # Оценка
        row_cells[3].text = '-'
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Допуск к защите
        admission = "Да"
        row_cells[4].text = admission
        row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_para.add_run('Руководитель подразделения: _________________')
    footer_run.font.size = Pt(12)

    temp_path = f"/tmp/{filename}"
    doc.save(temp_path)

    response = FileResponse(
        path=temp_path,
        filename=filename,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

    response.headers["X-Cleanup"] = temp_path
    return response